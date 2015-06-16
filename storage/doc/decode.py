#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx.table import Table
from docx.text.paragraph import Paragraph
from storage.doc.helpers import join_paragraphs, iter_block_items
from storage.doc.model import Node
from models.data_types import PrimitiveDataType, ComplexDataType
from models.resource import Resource
from models.method import Method
from models.permissions import Role, Access, Permission
from view.helpers import print_table


def populate_primitive_types(registry, table):
    for row in table.rows[1:]:
        name = row.cells[0].paragraphs[0].text.strip()
        description = join_paragraphs(row.cells[1].paragraphs)
        registry.add_type(PrimitiveDataType(name, description))


def populate_fields(data_type, table):
    for row in table.rows[1:]:
        name = row.cells[0].paragraphs[0].text.strip()
        datatypename = row.cells[2].paragraphs[0].text.strip()
        required = row.cells[1].paragraphs[0].text.strip() == u'+'
        description = join_paragraphs(row.cells[3].paragraphs)
        data_type.add_field(name, datatypename, description, required)


def get_complex_type_name(node):
    node_names = node.name.strip().split(u' ')
    if u'-' in node_names:
        node_names.remove(u'-')
    return node_names[0]


def is_complex_type_deprecated(node):
    node_names = node.name.strip().split(u' ')
    return u'DEPRECATED' in node_names


def get_complex_type_description(node):
    paragraphs = []
    for b in node.blocks:
        if isinstance(b, Paragraph):
            if b.text != node.name and not b.text.startswith(u'Таблица'):
                paragraphs.append(b)
    return join_paragraphs(paragraphs)


def get_resource_description(node):
    return get_complex_type_description(node)


def get_method_description(node):
    paragraphs = []
    result_header_index = None
    for idx, b in enumerate(node.blocks):
        if isinstance(b, Paragraph):
            if b.text.startswith(u'Результат'):
                result_header_index = idx
    for idx, b in enumerate(node.blocks):
        if isinstance(b, Paragraph):
            if b.text != node.name:
                continue
            if b.text.startswith(u'Таблица'):
                continue
            if b.text.startswith(u'Параметры'):
                continue
            if b.text.startswith(u'Результат'):
                continue
            if b.text.startswith(u'Уровень доступа для ролей'):
                continue
            if idx == result_header_index:
                continue
            paragraphs.append(b)
            result_header_index = None
    return join_paragraphs(paragraphs)


def parse_result_type(description):
    start_index = description.find(u'«')
    end_index = description.find(u'»')

    if description.find(u'Boolean.') != -1:
        return 'Boolean'
    if start_index != 1 and end_index != -1:
        if description.find(u'Массив') != -1:
            return 'Array.<' + description[start_index+1:end_index].split(' ')[-1] + '>'
        else:
            return description[start_index+1:end_index].split(' ')[-1]
    return None


def get_result_type_name(node):
    result_header_index = None
    for idx, b in enumerate(node.blocks):
        if isinstance(b, Paragraph):
            if b.text.startswith(u'Результат'):
                result_header_index = idx
    if result_header_index:
        return parse_result_type(node.blocks[result_header_index+1].text)
    else:
        return None


def get_tables(node):
    return filter(lambda x: isinstance(x, Table), node.blocks)


def _filter_methods_without_args(node):
    if get_complex_type_name(node) in [
        'Settings.show',
        'Settings.showAcceptanceByEmailSettings',
        'Settings.showFeatureStates',
        'Settings.showCallSettings',
        'Settings.showCobrowseSettings',
        'Settings.showEmployeeRemarkSettings',
        'Settings.showFileTransferSettings',
        'Settings.showReportByEmailSettings',
        'Settings.showTypingIndicatorSettings',
        'Stats.contacts'
    ]:
        return False
    else:
        return True


def get_complex_type(node):
    deprecated = is_complex_type_deprecated(node)
    name = get_complex_type_name(node)
    description = get_complex_type_description(node)
    data_type = ComplexDataType(name, description, deprecated)

    tables = get_tables(node)
    if len(tables) and _filter_methods_without_args(node):
        populate_fields(data_type, tables[0])

    return data_type


def parse_permissions_table(table):
    permissions = []
    for row in table.rows[1:]:
        role = Role.get_role(row.cells[0].paragraphs[0].text.strip())
        access = Access.get_access(row.cells[1].paragraphs[0].text.strip())
        description = row.cells[2].paragraphs[0].text.strip()

        permissions.append(Permission(role, access, description))

    return permissions


def get_permissions(node):
    heading_index = None
    for idx, block in enumerate(node.blocks):
        if isinstance(block, Paragraph):
            if block.text.find(u'ровень доступа для ролей') != -1:
                heading_index = idx
    permissions = []
    if heading_index:
        if isinstance(node.blocks[heading_index+1], Table):
            permissions = parse_permissions_table(node.blocks[heading_index+1])

    return permissions


def get_method(node):
    deprecated = is_complex_type_deprecated(node)
    name = parse_method_name(get_complex_type_name(node))
    description = get_method_description(node)
    result = get_result_type_name(node)
    method = Method(name, description, result, deprecated)

    tables = get_tables(node)
    if len(tables) and _filter_methods_without_args(node):
        populate_fields(method, tables[0])

    for perm in get_permissions(node):
        method.add_permission(perm)

    return method


def populate_complex_type(registry, node):
    registry.add_type(
        get_complex_type(node))

    for ch in node.children:
        populate_complex_type(registry, ch)


def get_heading_level(prev_level, block):
    if is_heading(block):
        return int(block.style.name[-1])
    else:
        return prev_level


def is_heading(block):
    return (isinstance(block, Paragraph) and
            (block.style.name.startswith(u'Заголовок') or
             block.style.name.startswith(u'Heading')))


def populate_nodes_from_doc(doc):
    root = Node(u'root')
    current_node = root
    current_level = 0

    for block in iter_block_items(doc):
        if is_heading(block):
            level = get_heading_level(current_level, block)
            if current_level >= level:
                for _ in range(0, current_level - level + 1):
                    current_node = current_node.parent
            node = Node(block.text, current_node)
            current_node = node
            current_level = level
        current_node.add_block(block)

    return root


def populate_types_from_node(root, reg):
    # Populate Primitive Data Types
    primitive_type_blocks = root.children[1].children[0].blocks
    for block in primitive_type_blocks:
        if isinstance(block, Table):
            populate_primitive_types(reg, block)

    # Populate Complex Data Types
    complex_type_nodes = root.children[1].children[1].children
    for node in complex_type_nodes:
        populate_complex_type(reg, node)


def parse_method_name(name):
    return name.split(u'.')[1]


def populate_resource(reg, node):
    res = Resource(get_complex_type_name(node),
                   get_resource_description(node))

    for ch in node.children:
        method = get_method(ch)
        res.add_method(method)

    reg.add_type(res)


def populate_resources_from_node(root, reg):
    for node in root.children[2].children:
        populate_resource(reg, node)
